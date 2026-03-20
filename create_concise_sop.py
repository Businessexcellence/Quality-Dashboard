#!/usr/bin/env python3
"""
Create a concise, practical SOP Word document for Quality Dashboard
Focus on essential working information only
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(240, 70, 22)  # Taggd orange
        heading.runs[0].font.size = Pt(20)
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(59, 130, 246)  # Blue
        heading.runs[0].font.size = Pt(16)
    return heading

def add_formula_box(doc, formula_name, formula, description):
    """Add a formula in a clean format"""
    p = doc.add_paragraph()
    p.add_run(f"• {formula_name}: ").bold = True
    p.add_run(f"{formula}")
    if description:
        p = doc.add_paragraph(f"  {description}", style='List Bullet 2')

def create_concise_sop():
    """Create concise SOP document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Quality Dashboard - Standard Operating Procedure', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(240, 70, 22)
    
    # Document Info
    doc.add_paragraph('Version 1.0 | Business Excellence Team', style='Normal')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 1: GETTING STARTED
    # ============================================================
    add_heading(doc, '1. Getting Started', level=1)
    
    add_heading(doc, 'Access & Login', level=2)
    doc.add_paragraph('• URL: https://3000-iiqzlm50w1twpgggr1553-82b888ba.sandbox.novita.ai')
    doc.add_paragraph('• Password: Excellence@2026')
    doc.add_paragraph('• Upload Excel file: Base File.xlsx')
    doc.add_paragraph()
    
    add_heading(doc, 'Dashboard Navigation', level=2)
    tabs = [
        'Home - Overview of key metrics',
        'Journey at Glance - Visual audit flow',
        'Account Summary - Account-wise performance',
        'Transactional Overview - Monthly trends',
        'Parameter Performance - Parameter-wise analysis',
        'Recruiter Insights - Recruiter performance',
        'Strategic Overview - Regional & head-wise data',
        'Project Summary - Project-wise metrics',
        'RCA & CAPA - Root cause analysis',
        'CSAT - Customer satisfaction',
        'Competitive Intelligence - Industry benchmarks',
        'SLA Overview - SLA compliance'
    ]
    for tab in tabs:
        doc.add_paragraph(f'• {tab}', style='List Bullet')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 2: KEY METRICS & FORMULAS
    # ============================================================
    add_heading(doc, '2. Key Metrics & Formulas', level=1)
    
    add_heading(doc, 'Core Accuracy Metrics', level=2)
    add_formula_box(doc, 
        'Overall Accuracy (%)', 
        '(Opportunity Pass / (Total Opportunities - NA)) × 100',
        'Primary quality metric. Target: ≥95% (Green), 90-94% (Amber), <90% (Red)')
    
    add_formula_box(doc,
        'Critical Accuracy (%)',
        '(Critical Pass / (Critical Count - Critical NA)) × 100',
        'Accuracy for critical parameters only')
    
    add_formula_box(doc,
        'Non-Critical Accuracy (%)',
        '(Non-Critical Pass / (Non-Critical Count - Non-Critical NA)) × 100',
        'Accuracy for non-critical parameters')
    doc.add_paragraph()
    
    add_heading(doc, 'Audit Coverage Metrics', level=2)
    add_formula_box(doc,
        'Sample Size (%)',
        '(Total Audits / Total Population) × 100',
        'Audit coverage. Target: ≥5%')
    
    add_formula_box(doc,
        'Error Rate (%)',
        '(Opportunity Fail / (Total Opportunities - NA)) × 100',
        'Error percentage. Target: ≤5%')
    doc.add_paragraph()
    
    add_heading(doc, 'CSAT Metrics', level=2)
    add_formula_box(doc,
        'Average CSAT Score',
        'Sum of all ratings / Total responses',
        'Scale: 1-7. Target: ≥6.0')
    
    add_formula_box(doc,
        'Top Box % (Satisfied)',
        '(Count of ratings 6-7 / Total responses) × 100',
        'Percentage giving ratings 6 or 7')
    
    add_formula_box(doc,
        'Response Rate (%)',
        '(Total Responses / Total Surveys Sent) × 100',
        'Survey completion rate')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 3: TAB-BY-TAB GUIDE
    # ============================================================
    add_heading(doc, '3. Tab-by-Tab Working Guide', level=1)
    
    # HOME TAB
    add_heading(doc, '3.1 Home Tab', level=2)
    doc.add_paragraph('Displays 6 KPI cards showing overall performance:')
    kpis = [
        'Overall Accuracy: Primary metric with color coding',
        'Critical Accuracy: Critical parameters only',
        'Non-Critical Accuracy: Non-critical parameters',
        'Sample Size: Audit coverage percentage',
        'Error Rate: Overall error percentage',
        'Total Audits: Count of completed audits'
    ]
    for kpi in kpis:
        doc.add_paragraph(f'• {kpi}', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Color Coding: ').bold = True
    p.add_run('Green (≥95%), Amber (90-94%), Red (<90%)')
    doc.add_paragraph()
    
    # JOURNEY AT GLANCE
    add_heading(doc, '3.2 Journey at Glance', level=2)
    doc.add_paragraph('Visual funnel showing audit progression:')
    doc.add_paragraph('• Total Population → Audited → Opportunity Count → Pass/Fail', style='List Bullet')
    doc.add_paragraph('• Filter by: Account Type, Region, Stage, Month', style='List Bullet')
    doc.add_paragraph('• Click stages to view detailed breakdown', style='List Bullet')
    doc.add_paragraph()
    
    # ACCOUNT SUMMARY
    add_heading(doc, '3.3 Account Summary', level=2)
    doc.add_paragraph('Account-wise performance table with:')
    doc.add_paragraph('• Accuracy %, Sample %, Error % for each account', style='List Bullet')
    doc.add_paragraph('• Total Population and Audit Count', style='List Bullet')
    doc.add_paragraph('• Pass, Fail, NA counts', style='List Bullet')
    doc.add_paragraph('• Search and filter by account name', style='List Bullet')
    doc.add_paragraph()
    
    # TRANSACTIONAL OVERVIEW
    add_heading(doc, '3.4 Transactional Overview', level=2)
    doc.add_paragraph('Monthly trend analysis:')
    doc.add_paragraph('• Month-wise chart: Accuracy % (line) + Audit Count (bars)', style='List Bullet')
    doc.add_paragraph('• Time series showing 8 months of data', style='List Bullet')
    doc.add_paragraph('• Filter by Region, Account, Stage', style='List Bullet')
    doc.add_paragraph('• Legend: 🔵 Accuracy Score (%) | 🟧 Audit Count', style='List Bullet')
    doc.add_paragraph()
    
    # PARAMETER PERFORMANCE
    add_heading(doc, '3.5 Parameter Performance', level=2)
    doc.add_paragraph('Parameter-wise accuracy breakdown:')
    doc.add_paragraph('• Table showing each parameter accuracy', style='List Bullet')
    doc.add_paragraph('• Critical vs Non-Critical classification', style='List Bullet')
    doc.add_paragraph('• Pass, Fail, NA counts per parameter', style='List Bullet')
    doc.add_paragraph('• Identify weak parameters for training', style='List Bullet')
    doc.add_paragraph()
    
    # RECRUITER INSIGHTS
    add_heading(doc, '3.6 Recruiter Insights', level=2)
    doc.add_paragraph('Individual recruiter performance:')
    doc.add_paragraph('• Accuracy % for each recruiter', style='List Bullet')
    doc.add_paragraph('• Audit count and error rate', style='List Bullet')
    doc.add_paragraph('• Filter by Region, Account, Stage', style='List Bullet')
    doc.add_paragraph('• Use for performance reviews and coaching', style='List Bullet')
    doc.add_paragraph()
    
    # STRATEGIC OVERVIEW
    add_heading(doc, '3.7 Strategic Overview', level=2)
    doc.add_paragraph('Regional and leadership view:')
    doc.add_paragraph('• Regional Head-wise performance table', style='List Bullet')
    doc.add_paragraph('• Practice Head-wise metrics', style='List Bullet')
    doc.add_paragraph('• India map with region-wise data', style='List Bullet')
    doc.add_paragraph('• Click regions for drill-down view', style='List Bullet')
    doc.add_paragraph()
    
    # RCA & CAPA
    add_heading(doc, '3.8 RCA & CAPA', level=2)
    doc.add_paragraph('Root Cause Analysis & Corrective Actions:')
    doc.add_paragraph('• Top 5 error categories', style='List Bullet')
    doc.add_paragraph('• Reason-wise failure breakdown', style='List Bullet')
    doc.add_paragraph('• Action items and owners', style='List Bullet')
    doc.add_paragraph('• Use for quality improvement initiatives', style='List Bullet')
    doc.add_paragraph()
    
    # CSAT
    add_heading(doc, '3.9 CSAT Tab', level=2)
    doc.add_paragraph('Customer satisfaction analysis:')
    doc.add_paragraph('• Top Box cards: Overall Score, Promoters %, Response Rate', style='List Bullet')
    doc.add_paragraph('• Rating distribution (1-7 scale)', style='List Bullet')
    doc.add_paragraph('• Account-wise CSAT scores', style='List Bullet')
    doc.add_paragraph('• Key Insights & Analysis section', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Current Findings:').bold = True
    doc.add_paragraph('• Response Rate: 61% (absolute numbers increasing)', style='List Bullet')
    doc.add_paragraph('• Overall Satisfaction: 5.13 (down from 5.95)', style='List Bullet')
    doc.add_paragraph('• Positives: End-to-end hiring, engagement, flexibility', style='List Bullet')
    doc.add_paragraph('• Focus Areas: Niche hiring, team quality, better profiles', style='List Bullet')
    doc.add_paragraph()
    
    # COMPETITIVE INTELLIGENCE
    add_heading(doc, '3.10 Competitive Intelligence', level=2)
    doc.add_paragraph('Industry benchmark comparison:')
    doc.add_paragraph('• Compare Taggd vs Industry standards', style='List Bullet')
    doc.add_paragraph('• Filter by: Global RPO, Indian RPO, Global Service Provider, Indian Service Provider', style='List Bullet')
    doc.add_paragraph('• Metrics: Quality Score, CSAT Score, TAT (days), Offer-to-Join %', style='List Bullet')
    doc.add_paragraph()
    
    # SLA OVERVIEW
    add_heading(doc, '3.11 SLA Overview', level=2)
    doc.add_paragraph('SLA compliance tracking:')
    doc.add_paragraph('• Total SLA Met vs Breached', style='List Bullet')
    doc.add_paragraph('• Average TAT (days)', style='List Bullet')
    doc.add_paragraph('• Stage-wise SLA performance', style='List Bullet')
    doc.add_paragraph('• Target: 80% SLA compliance', style='List Bullet')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 4: DATA UPLOAD & MANAGEMENT
    # ============================================================
    add_heading(doc, '4. Data Upload & Management', level=1)
    
    add_heading(doc, 'Uploading Data', level=2)
    doc.add_paragraph('1. Click "Upload Excel Data" button on any tab')
    doc.add_paragraph('2. Select Base File.xlsx from your computer')
    doc.add_paragraph('3. Wait for upload confirmation (usually 5-10 seconds)')
    doc.add_paragraph('4. Dashboard will refresh automatically with new data')
    doc.add_paragraph()
    
    add_heading(doc, 'Excel File Requirements', level=2)
    doc.add_paragraph('Required sheets in Base File.xlsx:')
    sheets = [
        'Sheet1: Main audit data',
        'CSAT: Customer satisfaction survey data',
        'SLA Data: SLA tracking information'
    ]
    for sheet in sheets:
        doc.add_paragraph(f'• {sheet}', style='List Bullet')
    doc.add_paragraph()
    
    add_heading(doc, 'Required Columns', level=2)
    columns = [
        'Audit Date, Account Name, Region, Stage',
        'Recruiter Name, Regional Head, Practice Head',
        'Parameter Name, Critical/Non-Critical, Result (Pass/Fail/NA)',
        'Total Population, Opportunity Count',
        'For CSAT: Question, Rating (1-7), Account, Date',
        'For SLA: Stage, TAT (days), SLA Status (Met/Breached)'
    ]
    for col in columns:
        doc.add_paragraph(f'• {col}', style='List Bullet')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 5: FILTERS & INTERACTIVITY
    # ============================================================
    add_heading(doc, '5. Using Filters', level=1)
    
    add_heading(doc, 'Global Filters (Available on Most Tabs)', level=2)
    doc.add_paragraph('• Account Type: Key Accounts, Non-Key Accounts, OPS, All')
    doc.add_paragraph('• Region: North, South, East, West, All')
    doc.add_paragraph('• Stage: Sourcing, Screening, Interview, All')
    doc.add_paragraph('• Month: Filter by specific month or All')
    doc.add_paragraph()
    
    add_heading(doc, 'How to Apply Filters', level=2)
    doc.add_paragraph('1. Click filter button (changes color when active)')
    doc.add_paragraph('2. Dashboard updates automatically')
    doc.add_paragraph('3. Multiple filters can be combined')
    doc.add_paragraph('4. Click "All" to reset specific filter')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 6: COLOR CODING REFERENCE
    # ============================================================
    add_heading(doc, '6. Color Coding Reference', level=1)
    
    add_heading(doc, 'Accuracy Colors', level=2)
    p = doc.add_paragraph()
    run = p.add_run('🟢 Green: ')
    run.bold = True
    run.font.color.rgb = RGBColor(22, 163, 74)
    p.add_run('≥95% (Excellent)')
    
    p = doc.add_paragraph()
    run = p.add_run('🟡 Amber: ')
    run.bold = True
    run.font.color.rgb = RGBColor(245, 158, 11)
    p.add_run('90-94% (Good, needs attention)')
    
    p = doc.add_paragraph()
    run = p.add_run('🔴 Red: ')
    run.bold = True
    run.font.color.rgb = RGBColor(239, 68, 68)
    p.add_run('<90% (Critical, immediate action required)')
    doc.add_paragraph()
    
    add_heading(doc, 'CSAT Rating Colors', level=2)
    p = doc.add_paragraph()
    run = p.add_run('🟢 Green: ')
    run.bold = True
    run.font.color.rgb = RGBColor(22, 163, 74)
    p.add_run('6-7 (Promoters/Satisfied)')
    
    p = doc.add_paragraph()
    run = p.add_run('🟡 Amber: ')
    run.bold = True
    run.font.color.rgb = RGBColor(245, 158, 11)
    p.add_run('4-5 (Neutral/Passives)')
    
    p = doc.add_paragraph()
    run = p.add_run('🔴 Red: ')
    run.bold = True
    run.font.color.rgb = RGBColor(239, 68, 68)
    p.add_run('1-3 (Detractors/Dissatisfied)')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 7: QUICK REFERENCE
    # ============================================================
    add_heading(doc, '7. Quick Reference Table', level=1)
    
    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Target'
    hdr_cells[2].text = 'Action If Below Target'
    
    # Data rows
    metrics_data = [
        ('Overall Accuracy', '≥95%', 'Review failed parameters, provide training'),
        ('Critical Accuracy', '≥98%', 'Immediate corrective action required'),
        ('Sample Size', '≥5%', 'Increase audit coverage'),
        ('Error Rate', '≤5%', 'Identify root causes, implement CAPA'),
        ('CSAT Score', '≥6.0', 'Review feedback, improve service delivery'),
        ('Top Box %', '≥70%', 'Focus on customer engagement initiatives'),
        ('Response Rate', '≥65%', 'Follow-up with customers, incentivize surveys'),
        ('SLA Compliance', '≥80%', 'Review process bottlenecks, optimize TAT')
    ]
    
    for metric, target, action in metrics_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = target
        row_cells[2].text = action
    
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 8: TROUBLESHOOTING
    # ============================================================
    add_heading(doc, '8. Common Issues & Solutions', level=1)
    
    issues = [
        ('Dashboard shows "No data available"', 'Upload Base File.xlsx using Upload button'),
        ('Filters not working', 'Refresh page (F5) and try again'),
        ('Charts not displaying', 'Clear browser cache, refresh page'),
        ('Upload fails', 'Check Excel file format and required sheets'),
        ('Incorrect calculations', 'Verify data in Base File.xlsx columns'),
        ('Month-wise chart overlapping', 'Already fixed - legend positioned above chart')
    ]
    
    for issue, solution in issues:
        p = doc.add_paragraph()
        p.add_run(f'Issue: {issue}').bold = True
        doc.add_paragraph(f'Solution: {solution}', style='List Bullet')
    
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 9: BEST PRACTICES
    # ============================================================
    add_heading(doc, '9. Best Practices', level=1)
    
    practices = [
        'Upload fresh data weekly for accurate trends',
        'Review Home tab KPIs daily for quick health check',
        'Use filters to drill down into specific issues',
        'Check CSAT tab monthly for customer feedback trends',
        'Review RCA & CAPA for quality improvement opportunities',
        'Monitor Transactional Overview for monthly patterns',
        'Use Recruiter Insights for coaching and training needs',
        'Compare with Competitive Intelligence for market positioning',
        'Track SLA Overview to ensure timely delivery',
        'Export data regularly for offline analysis'
    ]
    
    for practice in practices:
        doc.add_paragraph(f'• {practice}', style='List Bullet')
    
    doc.add_paragraph()
    
    # ============================================================
    # FOOTER
    # ============================================================
    doc.add_page_break()
    add_heading(doc, 'Document Information', level=1)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light List Accent 1'
    
    info_data = [
        ('Document Title', 'Quality Dashboard - Standard Operating Procedure'),
        ('Version', '1.0'),
        ('Department', 'Business Excellence Team'),
        ('Dashboard URL', 'https://3000-iiqzlm50w1twpgggr1553-82b888ba.sandbox.novita.ai'),
        ('Password', 'Excellence@2026')
    ]
    
    for i, (key, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = value
    
    # Save document
    output_path = '/home/user/webapp/public/downloads/Quality_Dashboard_SOP_Concise.docx'
    doc.save(output_path)
    print(f"✅ Concise SOP document created: {output_path}")
    print(f"📄 Document size: ~20-25 pages (vs previous ~100 pages)")
    print(f"🎯 Focus: Practical working guide with formulas, metrics, and procedures")
    
    return output_path

if __name__ == '__main__':
    create_concise_sop()
