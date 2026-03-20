#!/usr/bin/env python3
"""
Convert Markdown documentation to Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document(md_file, output_file):
    """Convert markdown to Word document"""
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = 'Quality Dashboard - Complete Documentation'
    doc.core_properties.author = 'Quality Dashboard Team'
    doc.core_properties.subject = 'Dashboard User Guide & Technical Documentation'
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    lines = content.split('\n')
    
    # Process lines
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at start
        if not line.strip() and i == 0:
            i += 1
            continue
        
        # Heading 1 (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.style.font.color.rgb = RGBColor(240, 70, 22)  # Taggd orange
        
        # Heading 2 (## )
        elif line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.style.font.color.rgb = RGBColor(59, 130, 246)  # Blue
        
        # Heading 3 (### )
        elif line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
        
        # Heading 4 (#### )
        elif line.startswith('#### '):
            text = line[5:].strip()
            doc.add_heading(text, level=4)
        
        # Horizontal rule (---)
        elif line.strip() == '---':
            doc.add_paragraph()  # Add spacing
        
        # Code blocks (```)
        elif line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Quote'
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)
                
                # Set font to monospace
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Bullet lists (- or * at start)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            # Remove markdown bold/italic
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\. ', line.strip()):
            text = re.sub(r'^\d+\. ', '', line.strip())
            # Remove markdown bold/italic
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
        
        # Regular paragraph
        elif line.strip():
            text = line.strip()
            
            # Skip markdown-only lines
            if text in ['---', '===']:
                i += 1
                continue
            
            # Remove markdown bold/italic for regular text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            
            # Add paragraph
            if not text.startswith('#'):
                p = doc.add_paragraph(text)
                
                # Format bold text
                if '**' in line:
                    # Re-process with bold
                    p.clear()
                    parts = re.split(r'(\*\*.+?\*\*)', line.strip())
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)
        
        i += 1
    
    # Add page numbers
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Quality Dashboard Documentation - Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")

if __name__ == '__main__':
    md_file = 'Quality_Dashboard_Complete_Documentation.md'
    output_file = 'Quality_Dashboard_Documentation.docx'
    
    print(f"📄 Converting {md_file} to Word format...")
    create_word_document(md_file, output_file)
    print(f"✅ Conversion complete!")
    print(f"📁 File location: /home/user/webapp/{output_file}")
