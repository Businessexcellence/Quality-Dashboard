#!/usr/bin/env python3
"""
Create a concise 10-page Word document covering complete dashboard working
Focus: All tabs, key metrics, essential formulas only
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(240, 70, 22)
        heading.runs[0].font.size = Pt(18)
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(59, 130, 246)
        heading.runs[0].font.size = Pt(14)
    return heading

def create_compact_sop():
    """Create 10-page concise SOP"""
    doc = Document()
    
    # Set narrow margins for more content
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    title = doc.add_heading('Quality Dashboard - Complete Working Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(240, 70, 22)
    title.runs[0].font.size = Pt(22)
    
    # Subtitle
    p = doc.add_paragraph('Standard Operating Procedure | Version 1.0')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 1: ACCESS & UPLOAD
    # ============================================================
    add_heading(doc, '1. Access & Data Upload', level=1)
    
    doc.add_paragraph('• URL: https://3000-iiqzlm50w1twpgggr1553-82b888ba.sandbox.novita.ai')
    doc.add_paragraph('• Password: Excellence@2026')
    doc.add_paragraph('• Upload: Click "Upload Excel Data" → Select Base File.xlsx → Wait for confirmation')
    doc.add_paragraph('• Required Sheets: Sheet1 (main data), CSAT (survey data), SLA Data (SLA tracking)')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 2: KEY FORMULAS
    # ============================================================
    add_heading(doc, '2. Key Formulas', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Overall Accuracy = ').bold = True
    p.add_run('(Pass / (Total - NA)) × 100  ')
    p.add_run('[Target: ≥95%]').italic = True
    
    p = doc.add_paragraph()
    p.add_run('Critical Accuracy = ').bold = True
    p.add_run('(Critical Pass / (Critical Count - NA)) × 100  ')
    p.add_run('[Target: ≥98%]').italic = True
    
    p = doc.add_paragraph()
    p.add_run('Sample Size % = ').bold = True
    p.add_run('(Audits / Population) × 100  ')
    p.add_run('[Target: ≥5%]').italic = True
    
    p = doc.add_paragraph()
    p.add_run('Error Rate % = ').bold = True
    p.add_run('(Fail / (Total - NA)) × 100  ')
    p.add_run('[Target: ≤5%]').italic = True
    
    p = doc.add_paragraph()
    p.add_run('CSAT Score = ').bold = True
    p.add_run('Sum of Ratings / Total Responses  ')
    p.add_run('[Target: ≥6.0 on 1-7 scale]').italic = True
    
    p = doc.add_paragraph()
    p.add_run('Top Box % = ').bold = True
    p.add_run('(Ratings 6-7 / Total) × 100  ')
    p.add_run('[Target: ≥70%]').italic = True
    
    p = doc.add_paragraph()
    p.add_run('Color Coding: ').bold = True
    run = p.add_run('Green ≥95%')
    run.font.color.rgb = RGBColor(22, 163, 74)
    p.add_run(' | ')
    run = p.add_run('Amber 90-94%')
    run.font.color.rgb = RGBColor(245, 158, 11)
    p.add_run(' | ')
    run = p.add_run('Red <90%')
    run.font.color.rgb = RGBColor(239, 68, 68)
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 3: ALL TABS WORKING
    # ============================================================
    add_heading(doc, '3. Dashboard Tabs - Complete Working', level=1)
    
    # HOME TAB
    add_heading(doc, '3.1 Home Tab', level=2)
    doc.add_paragraph('Shows 6 KPI cards: Overall Accuracy, Critical Accuracy, Non-Critical Accuracy, Sample Size, Error Rate, Total Audits. Color-coded based on performance. Use for daily health check.')
    
    # JOURNEY AT GLANCE
    add_heading(doc, '3.2 Journey at Glance', level=2)
    doc.add_paragraph('Visual funnel: Population → Audited → Opportunities → Pass/Fail. Filters: Account Type, Region, Stage, Month. Click stages for breakdown. Shows audit flow from start to end.')
    
    # ACCOUNT SUMMARY
    add_heading(doc, '3.3 Account Summary', level=2)
    doc.add_paragraph('Table with all accounts showing: Accuracy %, Sample %, Error %, Population, Audits, Pass, Fail, NA. Search by account name. Filter by account type. Use for account-wise performance review.')
    
    # TRANSACTIONAL OVERVIEW
    add_heading(doc, '3.4 Transactional Overview', level=2)
    doc.add_paragraph('Month-wise chart with Accuracy % (blue line) and Audit Count (orange bars). Shows 8 months trend. Filters: Region, Account, Stage. Legend positioned above chart. Use for monthly trend analysis.')
    
    # PARAMETER PERFORMANCE
    add_heading(doc, '3.5 Parameter Performance', level=2)
    doc.add_paragraph('Table showing accuracy for each parameter. Critical vs Non-Critical classification. Pass, Fail, NA counts. Identify weak parameters for training. Filter by region and account.')
    
    # RECRUITER INSIGHTS
    add_heading(doc, '3.6 Recruiter Insights', level=2)
    doc.add_paragraph('Individual recruiter performance: Accuracy %, Audits, Errors. Filter by Region, Account, Stage. Use for coaching and performance reviews. Sort by any column.')
    
    # STRATEGIC OVERVIEW
    add_heading(doc, '3.7 Strategic Overview', level=2)
    doc.add_paragraph('Regional Head and Practice Head performance tables. India map showing region-wise data (click regions for drill-down). Filter by account type. Use for leadership review.')
    
    # PROJECT SUMMARY
    add_heading(doc, '3.8 Project Summary', level=2)
    doc.add_paragraph('Project-wise metrics: Accuracy %, Audits, Population. Filter by region and month. Use for project-level tracking and resource allocation planning.')
    
    # RCA & CAPA
    add_heading(doc, '3.9 RCA & CAPA', level=2)
    doc.add_paragraph('Root Cause Analysis: Top 5 error categories with counts. Reason-wise failure breakdown. Action items with owners. Use for quality improvement initiatives and corrective actions.')
    
    # CSAT
    add_heading(doc, '3.10 CSAT Tab', level=2)
    doc.add_paragraph('Customer Satisfaction: Top Box cards (Overall Score, Promoters %, Response Rate). Rating distribution (1-7 scale). Account-wise CSAT. Key Insights: Response Rate 61%, Score 5.13 (down from 5.95), Focus: Niche hiring, team quality.')
    
    # COMPETITIVE INTELLIGENCE
    add_heading(doc, '3.11 Competitive Intelligence', level=2)
    doc.add_paragraph('Industry benchmarks comparison. Filters: Global RPO, Indian RPO, Global Service Provider, Indian Service Provider. Metrics: Quality Score, CSAT, TAT, Offer-to-Join %. Compare Taggd vs industry standards.')
    
    # SLA OVERVIEW
    add_heading(doc, '3.12 SLA Overview', level=2)
    doc.add_paragraph('SLA compliance tracking: Total Met vs Breached, Average TAT (days), Stage-wise performance. Target: 80% compliance. Filter by account and region. Use for process optimization.')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 4: FILTERS USAGE
    # ============================================================
    add_heading(doc, '4. Using Filters', level=1)
    
    doc.add_paragraph('• Account Type: Key Accounts, Non-Key Accounts, OPS, All')
    doc.add_paragraph('• Region: North, South, East, West, All')
    doc.add_paragraph('• Stage: Sourcing, Screening, Interview, All')
    doc.add_paragraph('• Month: Individual months or All')
    doc.add_paragraph('• How: Click filter button (turns active), combine multiple filters, click "All" to reset')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 5: QUICK REFERENCE TABLE
    # ============================================================
    add_heading(doc, '5. Quick Reference', level=1)
    
    # Compact table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Target'
    hdr_cells[2].text = 'Action'
    
    # Data
    metrics = [
        ('Overall Accuracy', '≥95%', 'Review parameters, provide training'),
        ('Critical Accuracy', '≥98%', 'Immediate corrective action'),
        ('Sample Size', '≥5%', 'Increase audit coverage'),
        ('Error Rate', '≤5%', 'Identify root causes, CAPA'),
        ('CSAT Score', '≥6.0', 'Review feedback, improve service'),
        ('Top Box %', '≥70%', 'Customer engagement initiatives'),
        ('Response Rate', '≥65%', 'Follow-up, incentivize surveys'),
        ('SLA Compliance', '≥80%', 'Review bottlenecks, optimize TAT')
    ]
    
    for metric, target, action in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = target
        row[2].text = action
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 6: EXCEL FILE STRUCTURE
    # ============================================================
    add_heading(doc, '6. Excel File Structure', level=1)
    
    add_heading(doc, 'Sheet1 (Main Data) - Required Columns:', level=2)
    doc.add_paragraph('Audit Date, Account Name, Region, Stage, Recruiter Name, Regional Head, Practice Head, Parameter Name, Critical/Non-Critical, Result (Pass/Fail/NA), Total Population, Opportunity Count')
    
    add_heading(doc, 'CSAT Sheet - Required Columns:', level=2)
    doc.add_paragraph('Question, Rating (1-7 scale), Account, Date, Response ID')
    
    add_heading(doc, 'SLA Data Sheet - Required Columns:', level=2)
    doc.add_paragraph('Stage, TAT (days), SLA Status (Met/Breached), Account, Date')
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 7: TROUBLESHOOTING
    # ============================================================
    add_heading(doc, '7. Troubleshooting', level=1)
    
    issues = [
        ('No data shown', 'Upload Base File.xlsx via Upload button'),
        ('Filters not working', 'Refresh page (F5), reapply filters'),
        ('Charts blank', 'Clear browser cache, reload'),
        ('Upload fails', 'Check Excel format, required sheets'),
        ('Wrong calculations', 'Verify Excel column data')
    ]
    
    for issue, solution in issues:
        p = doc.add_paragraph()
        p.add_run(f'{issue}: ').bold = True
        p.add_run(solution)
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 8: BEST PRACTICES
    # ============================================================
    add_heading(doc, '8. Best Practices', level=1)
    
    doc.add_paragraph('• Daily: Check Home tab KPIs for health status')
    doc.add_paragraph('• Weekly: Upload fresh data, review trends in Transactional Overview')
    doc.add_paragraph('• Monthly: Analyze CSAT feedback, review RCA & CAPA items')
    doc.add_paragraph('• Use filters to drill down into specific issues')
    doc.add_paragraph('• Monitor recruiter performance for coaching needs')
    doc.add_paragraph('• Compare with industry benchmarks quarterly')
    doc.add_paragraph('• Track SLA compliance for process optimization')
    doc.add_paragraph()
    
    # ============================================================
    # FOOTER
    # ============================================================
    doc.add_paragraph('_' * 80)
    p = doc.add_paragraph()
    p.add_run('Dashboard URL: ').bold = True
    p.add_run('https://3000-iiqzlm50w1twpgggr1553-82b888ba.sandbox.novita.ai')
    
    p = doc.add_paragraph()
    p.add_run('Password: ').bold = True
    p.add_run('Excellence@2026')
    
    p = doc.add_paragraph()
    p.add_run('Support: ').bold = True
    p.add_run('Business Excellence Team')
    
    # Save
    output_path = '/home/user/webapp/public/downloads/Quality_Dashboard_SOP_10Pages.docx'
    doc.save(output_path)
    print(f"✅ 10-page SOP created: {output_path}")
    
    return output_path

if __name__ == '__main__':
    create_compact_sop()
