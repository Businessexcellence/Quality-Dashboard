#!/usr/bin/env python3
"""
Create a detailed 10-page Word document covering ALL dashboard tabs
Every card, every view, every calculation explained in detail
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(240, 70, 22)
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.bold = True
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(59, 130, 246)
        heading.runs[0].font.size = Pt(13)
        heading.runs[0].font.bold = True
    elif level == 3:
        heading.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        heading.runs[0].font.size = Pt(11)
        heading.runs[0].font.bold = True
    return heading

def add_formula(doc, name, formula, variables, example, color_coding=None):
    """Add detailed formula explanation"""
    p = doc.add_paragraph()
    p.add_run(name + ': ').bold = True
    p.paragraph_format.space_after = Pt(3)
    
    p = doc.add_paragraph(f'Formula: {formula}', style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    
    if variables:
        p = doc.add_paragraph(f'Variables: {variables}', style='List Bullet 2')
        p.paragraph_format.space_after = Pt(2)
    
    if example:
        p = doc.add_paragraph(f'Example: {example}', style='List Bullet 2')
        p.paragraph_format.space_after = Pt(2)
    
    if color_coding:
        p = doc.add_paragraph(f'Color Coding: {color_coding}', style='List Bullet 2')
        p.paragraph_format.space_after = Pt(4)

def create_detailed_10page_sop():
    """Create detailed 10-page SOP"""
    doc = Document()
    
    # Set margins for optimal content
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    
    # Title Page
    title = doc.add_heading('Quality Dashboard', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(240, 70, 22)
    title.runs[0].font.size = Pt(24)
    
    subtitle = doc.add_paragraph('Complete Working Guide & Standard Operating Procedure')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('Version 1.0 | Business Excellence Team')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    
    # Access Info Box
    p = doc.add_paragraph()
    p.add_run('Dashboard Access Information').bold = True
    p.add_run('\nURL: https://3000-iiqzlm50w1twpgggr1553-82b888ba.sandbox.novita.ai')
    p.add_run('\nPassword: Excellence@2026')
    p.add_run('\nData File: Base File.xlsx (Upload via "Upload Excel Data" button)')
    doc.add_paragraph()
    
    # ============================================================
    # PAGE 1-2: CORE FORMULAS & CALCULATIONS
    # ============================================================
    add_heading(doc, '1. Core Formulas & Calculations', level=1)
    
    add_heading(doc, '1.1 Accuracy Metrics', level=2)
    
    add_formula(doc,
        'Overall Accuracy (%)',
        '(Total Opportunity Pass / (Total Opportunity Count - Total Opportunity NA)) × 100',
        'Pass = audits passed, Count = total opportunities, NA = not applicable items',
        'Pass=2500, Count=2847, NA=150 → (2500/(2847-150))×100 = 92.7%',
        'Green ≥95%, Amber 90-94%, Red <90%')
    
    add_formula(doc,
        'Critical Accuracy (%)',
        '(Critical Opportunity Pass / (Critical Opportunity Count - Critical NA)) × 100',
        'Only includes parameters marked as "Critical"',
        'Critical Pass=1200, Critical Count=1250, NA=20 → (1200/(1250-20))×100 = 97.6%',
        'Green ≥98%, Amber 95-97%, Red <95%')
    
    add_formula(doc,
        'Non-Critical Accuracy (%)',
        '(Non-Critical Pass / (Non-Critical Count - Non-Critical NA)) × 100',
        'Only includes parameters marked as "Non-Critical"',
        'Non-Critical Pass=1300, Count=1597, NA=130 → (1300/(1597-130))×100 = 88.6%',
        'Green ≥90%, Amber 85-89%, Red <85%')
    
    add_heading(doc, '1.2 Coverage Metrics', level=2)
    
    add_formula(doc,
        'Sample Size (%)',
        '(Total Audits Conducted / Total Population) × 100',
        'Audits = number of audits done, Population = total available records',
        'Audits=7880, Population=150000 → (7880/150000)×100 = 5.25%',
        'Green ≥5%, Amber 3-4.9%, Red <3%')
    
    add_formula(doc,
        'Error Rate (%)',
        '(Total Opportunity Fail / (Total Opportunity Count - Total NA)) × 100',
        'Fail = failed opportunities, inverse of overall accuracy',
        'Fail=197, Count=2847, NA=150 → (197/(2847-150))×100 = 7.3%',
        'Green ≤5%, Amber 5.1-10%, Red >10%')
    
    add_heading(doc, '1.3 CSAT Metrics', level=2)
    
    add_formula(doc,
        'Average CSAT Score',
        'Sum of All Ratings / Total Number of Responses',
        'Rating scale: 1-7 (1=Very Dissatisfied, 7=Very Satisfied)',
        'Sum=308, Responses=60 → 308/60 = 5.13',
        'Green ≥6.0, Amber 5.0-5.9, Red <5.0')
    
    add_formula(doc,
        'Top Box % (Promoters)',
        '(Count of Ratings 6 or 7 / Total Responses) × 100',
        'Ratings 6-7 = Satisfied customers (Promoters)',
        'Ratings 6-7 = 35, Total=60 → (35/60)×100 = 58.3%',
        'Green ≥70%, Amber 60-69%, Red <60%')
    
    add_formula(doc,
        'Middle Box % (Passives)',
        '(Count of Ratings 4 or 5 / Total Responses) × 100',
        'Ratings 4-5 = Neutral customers (Passives)',
        'Ratings 4-5 = 18, Total=60 → (18/60)×100 = 30%',
        None)
    
    add_formula(doc,
        'Bottom Box % (Detractors)',
        '(Count of Ratings 1, 2 or 3 / Total Responses) × 100',
        'Ratings 1-3 = Dissatisfied customers (Detractors)',
        'Ratings 1-3 = 7, Total=60 → (7/60)×100 = 11.7%',
        None)
    
    add_formula(doc,
        'Response Rate (%)',
        '(Total Survey Responses Received / Total Surveys Sent) × 100',
        'Measures survey completion rate',
        'Responses=60, Sent=98 → (60/98)×100 = 61.2%',
        'Green ≥65%, Amber 50-64%, Red <50%')
    
    # ============================================================
    # PAGE 3-4: HOME TAB & JOURNEY AT GLANCE
    # ============================================================
    doc.add_page_break()
    add_heading(doc, '2. Home Tab - Dashboard Overview', level=1)
    
    doc.add_paragraph('The Home tab displays 6 primary KPI cards that provide a quick health check of overall quality performance. All cards are color-coded for instant visual assessment.')
    
    add_heading(doc, '2.1 KPI Cards Breakdown', level=2)
    
    # Card 1
    add_heading(doc, 'Card 1: Overall Accuracy', level=3)
    doc.add_paragraph('• Calculation: (Pass / (Total - NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Primary quality metric showing overall audit pass rate', style='List Bullet')
    doc.add_paragraph('• Data Source: Aggregates all audit results from Sheet1', style='List Bullet')
    doc.add_paragraph('• Icon: ✓ (checkmark)', style='List Bullet')
    doc.add_paragraph('• Target: ≥95% (Green), 90-94% (Amber), <90% (Red)', style='List Bullet')
    
    # Card 2
    add_heading(doc, 'Card 2: Critical Accuracy', level=3)
    doc.add_paragraph('• Calculation: (Critical Pass / (Critical Count - Critical NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Tracks accuracy for critical parameters only', style='List Bullet')
    doc.add_paragraph('• Data Source: Filters Sheet1 for "Critical" parameter type', style='List Bullet')
    doc.add_paragraph('• Icon: ⚠ (alert)', style='List Bullet')
    doc.add_paragraph('• Target: ≥98% (critical parameters have stricter targets)', style='List Bullet')
    
    # Card 3
    add_heading(doc, 'Card 3: Non-Critical Accuracy', level=3)
    doc.add_paragraph('• Calculation: (Non-Critical Pass / (Non-Critical Count - Non-Critical NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Tracks accuracy for non-critical parameters', style='List Bullet')
    doc.add_paragraph('• Data Source: Filters Sheet1 for "Non-Critical" parameter type', style='List Bullet')
    doc.add_paragraph('• Icon: ℹ (info)', style='List Bullet')
    doc.add_paragraph('• Target: ≥90% (slightly lower than overall target)', style='List Bullet')
    
    # Card 4
    add_heading(doc, 'Card 4: Sample Size %', level=3)
    doc.add_paragraph('• Calculation: (Total Audits / Total Population) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Measures audit coverage percentage', style='List Bullet')
    doc.add_paragraph('• Data Source: Counts unique audit IDs vs total population count', style='List Bullet')
    doc.add_paragraph('• Icon: 📊 (chart)', style='List Bullet')
    doc.add_paragraph('• Target: ≥5% (industry standard for audit sampling)', style='List Bullet')
    
    # Card 5
    add_heading(doc, 'Card 5: Error Rate %', level=3)
    doc.add_paragraph('• Calculation: (Fail / (Total - NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Percentage of failed audits (inverse of accuracy)', style='List Bullet')
    doc.add_paragraph('• Data Source: Counts all "Fail" results from Sheet1', style='List Bullet')
    doc.add_paragraph('• Icon: ✗ (cross)', style='List Bullet')
    doc.add_paragraph('• Target: ≤5% (lower is better, opposite of accuracy)', style='List Bullet')
    
    # Card 6
    add_heading(doc, 'Card 6: Total Audits', level=3)
    doc.add_paragraph('• Calculation: COUNT(DISTINCT Audit ID)', style='List Bullet')
    doc.add_paragraph('• Purpose: Total number of audits conducted in selected period', style='List Bullet')
    doc.add_paragraph('• Data Source: Unique count of audit records from Sheet1', style='List Bullet')
    doc.add_paragraph('• Icon: # (hash)', style='List Bullet')
    doc.add_paragraph('• Display: Absolute number (no percentage or color coding)', style='List Bullet')
    
    add_heading(doc, '2.2 How to Use Home Tab', level=2)
    doc.add_paragraph('1. Open dashboard and check all 6 cards for quick health status', style='List Number')
    doc.add_paragraph('2. Red cards indicate immediate attention required', style='List Number')
    doc.add_paragraph('3. Amber cards suggest monitoring and preventive action', style='List Number')
    doc.add_paragraph('4. Green cards show healthy performance', style='List Number')
    doc.add_paragraph('5. Use this tab daily for morning health check', style='List Number')
    doc.add_paragraph()
    
    # ============================================================
    # JOURNEY AT GLANCE
    # ============================================================
    add_heading(doc, '3. Journey at Glance Tab', level=1)
    
    doc.add_paragraph('This tab provides a visual funnel showing the audit journey from total population to final pass/fail results. It helps identify drop-off points and conversion rates at each stage.')
    
    add_heading(doc, '3.1 Funnel Stages', level=2)
    
    add_heading(doc, 'Stage 1: Total Population', level=3)
    doc.add_paragraph('• Calculation: SUM(Total Population) from all records', style='List Bullet')
    doc.add_paragraph('• Meaning: Total number of records available for auditing', style='List Bullet')
    doc.add_paragraph('• Example: 150,000 total candidates/records in database', style='List Bullet')
    
    add_heading(doc, 'Stage 2: Audited Records', level=3)
    doc.add_paragraph('• Calculation: COUNT(Distinct Audit Records)', style='List Bullet')
    doc.add_paragraph('• Meaning: Number of records actually audited', style='List Bullet')
    doc.add_paragraph('• Example: 7,880 audits conducted', style='List Bullet')
    doc.add_paragraph('• Conversion: (7,880 / 150,000) × 100 = 5.25% coverage', style='List Bullet')
    
    add_heading(doc, 'Stage 3: Opportunity Count', level=3)
    doc.add_paragraph('• Calculation: SUM(Opportunity Count) from audited records', style='List Bullet')
    doc.add_paragraph('• Meaning: Total checkpoints/parameters evaluated in audits', style='List Bullet')
    doc.add_paragraph('• Example: 2,847 opportunities (7,880 audits × avg parameters per audit)', style='List Bullet')
    
    add_heading(doc, 'Stage 4: Pass Count', level=3)
    doc.add_paragraph('• Calculation: COUNT(Result = "Pass")', style='List Bullet')
    doc.add_paragraph('• Meaning: Number of opportunities that passed audit', style='List Bullet')
    doc.add_paragraph('• Example: 2,500 opportunities passed', style='List Bullet')
    
    add_heading(doc, 'Stage 5: Fail Count', level=3)
    doc.add_paragraph('• Calculation: COUNT(Result = "Fail")', style='List Bullet')
    doc.add_paragraph('• Meaning: Number of opportunities that failed audit', style='List Bullet')
    doc.add_paragraph('• Example: 197 opportunities failed', style='List Bullet')
    
    add_heading(doc, 'Stage 6: NA Count', level=3)
    doc.add_paragraph('• Calculation: COUNT(Result = "NA")', style='List Bullet')
    doc.add_paragraph('• Meaning: Not applicable/not evaluated opportunities', style='List Bullet')
    doc.add_paragraph('• Example: 150 opportunities marked NA', style='List Bullet')
    
    add_heading(doc, '3.2 Available Filters', level=2)
    doc.add_paragraph('• Account Type: Key Accounts, Non-Key Accounts, OPS, All', style='List Bullet')
    doc.add_paragraph('• Region: North, South, East, West, All', style='List Bullet')
    doc.add_paragraph('• Stage: Sourcing, Screening, Interview, Onboarding, All', style='List Bullet')
    doc.add_paragraph('• Month: January through December, All', style='List Bullet')
    
    add_heading(doc, '3.3 Interactive Features', level=2)
    doc.add_paragraph('• Click any funnel stage to see detailed breakdown', style='List Bullet')
    doc.add_paragraph('• Hover over stages to see exact numbers and percentages', style='List Bullet')
    doc.add_paragraph('• Apply multiple filters simultaneously for drill-down analysis', style='List Bullet')
    doc.add_paragraph('• Export data by right-clicking on the funnel', style='List Bullet')
    
    # ============================================================
    # PAGE 5: ACCOUNT SUMMARY & TRANSACTIONAL OVERVIEW
    # ============================================================
    doc.add_page_break()
    add_heading(doc, '4. Account Summary Tab', level=1)
    
    doc.add_paragraph('Displays a comprehensive table showing performance metrics for each client account. Essential for account-wise performance review and identifying accounts needing attention.')
    
    add_heading(doc, '4.1 Table Columns Explained', level=2)
    
    add_heading(doc, 'Column 1: Account Name', level=3)
    doc.add_paragraph('• Source: Account Name from Sheet1', style='List Bullet')
    doc.add_paragraph('• Display: Client/account names (e.g., "Tech Corp", "Finance Inc")', style='List Bullet')
    doc.add_paragraph('• Sortable: Yes (alphabetically)', style='List Bullet')
    
    add_heading(doc, 'Column 2: Accuracy %', level=3)
    doc.add_paragraph('• Calculation: (Account Pass / (Account Total - Account NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Overall accuracy for this specific account', style='List Bullet')
    doc.add_paragraph('• Color Coded: Green ≥95%, Amber 90-94%, Red <90%', style='List Bullet')
    
    add_heading(doc, 'Column 3: Sample %', level=3)
    doc.add_paragraph('• Calculation: (Account Audits / Account Population) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Audit coverage for this account', style='List Bullet')
    doc.add_paragraph('• Target: ≥5%', style='List Bullet')
    
    add_heading(doc, 'Column 4: Error %', level=3)
    doc.add_paragraph('• Calculation: (Account Fail / (Account Total - Account NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Error rate specific to this account', style='List Bullet')
    doc.add_paragraph('• Target: ≤5%', style='List Bullet')
    
    add_heading(doc, 'Column 5: Total Population', level=3)
    doc.add_paragraph('• Calculation: SUM(Population) for account', style='List Bullet')
    doc.add_paragraph('• Purpose: Total records available for this account', style='List Bullet')
    doc.add_paragraph('• Display: Absolute number', style='List Bullet')
    
    add_heading(doc, 'Column 6: Audit Count', level=3)
    doc.add_paragraph('• Calculation: COUNT(Audits) for account', style='List Bullet')
    doc.add_paragraph('• Purpose: Number of audits conducted for this account', style='List Bullet')
    doc.add_paragraph('• Display: Absolute number', style='List Bullet')
    
    add_heading(doc, 'Columns 7-9: Pass, Fail, NA', level=3)
    doc.add_paragraph('• Calculation: COUNT(Result) grouped by Pass/Fail/NA', style='List Bullet')
    doc.add_paragraph('• Purpose: Breakdown of audit results', style='List Bullet')
    doc.add_paragraph('• Display: Absolute numbers', style='List Bullet')
    
    add_heading(doc, '4.2 Features', level=2)
    doc.add_paragraph('• Search Bar: Type account name for quick lookup', style='List Bullet')
    doc.add_paragraph('• Sort Function: Click column headers to sort', style='List Bullet')
    doc.add_paragraph('• Filter Options: Account Type, Region filters available', style='List Bullet')
    doc.add_paragraph('• Click Account: Opens detailed account drill-down view', style='List Bullet')
    
    # ============================================================
    # TRANSACTIONAL OVERVIEW
    # ============================================================
    add_heading(doc, '5. Transactional Overview Tab', level=1)
    
    doc.add_paragraph('Shows month-wise trends with a dual-axis chart displaying Accuracy % (line) and Audit Count (bars). Critical for identifying seasonal patterns and trends.')
    
    add_heading(doc, '5.1 Month Wise Chart Components', level=2)
    
    add_heading(doc, 'Primary Axis: Accuracy Score % (Blue Line)', level=3)
    doc.add_paragraph('• Calculation: Monthly (Pass / (Total - NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Display: Line graph with data points', style='List Bullet')
    doc.add_paragraph('• Color: Blue (#3b82f6)', style='List Bullet')
    doc.add_paragraph('• Scale: 0-100% on left Y-axis', style='List Bullet')
    doc.add_paragraph('• Data Labels: Shown above each point (e.g., "99.4%")', style='List Bullet')
    
    add_heading(doc, 'Secondary Axis: Audit Count (Orange Bars)', level=3)
    doc.add_paragraph('• Calculation: COUNT(Audits) per month', style='List Bullet')
    doc.add_paragraph('• Display: Vertical bars', style='List Bullet')
    doc.add_paragraph('• Color: Orange (#fb923c)', style='List Bullet')
    doc.add_paragraph('• Scale: Dynamic based on max count on right Y-axis', style='List Bullet')
    doc.add_paragraph('• Data Labels: Shown inside bars (e.g., "3,083")', style='List Bullet')
    
    add_heading(doc, 'X-Axis: Months', level=3)
    doc.add_paragraph('• Display: Last 8 months (e.g., Aug, Sep, Oct, Nov, Dec, Jan, Feb, Mar)', style='List Bullet')
    doc.add_paragraph('• Format: 3-letter month abbreviation', style='List Bullet')
    
    add_heading(doc, 'Legend', level=3)
    doc.add_paragraph('• Position: Above chart (external to plot area)', style='List Bullet')
    doc.add_paragraph('• Items: 🔵 Accuracy Score (%) | 🟧 Audit Count', style='List Bullet')
    doc.add_paragraph('• Purpose: No overlap with data points', style='List Bullet')
    
    add_heading(doc, '5.2 How to Read the Chart', level=2)
    doc.add_paragraph('1. Blue line shows quality trend - upward is good', style='List Number')
    doc.add_paragraph('2. Orange bars show audit volume - higher means more coverage', style='List Number')
    doc.add_paragraph('3. Compare trends: Are low-accuracy months also low-volume?', style='List Number')
    doc.add_paragraph('4. Hover over points for detailed tooltips with exact values', style='List Number')
    
    add_heading(doc, '5.3 Available Filters', level=2)
    doc.add_paragraph('• Region Filter: Filter chart by specific region', style='List Bullet')
    doc.add_paragraph('• Account Filter: Filter by specific account or account type', style='List Bullet')
    doc.add_paragraph('• Stage Filter: Filter by recruitment stage', style='List Bullet')
    doc.add_paragraph('• Multiple Filters: Combine filters for detailed trend analysis', style='List Bullet')
    
    # ============================================================
    # PAGE 6-7: PARAMETER, RECRUITER, STRATEGIC
    # ============================================================
    doc.add_page_break()
    add_heading(doc, '6. Parameter Performance Tab', level=1)
    
    doc.add_paragraph('Breaks down accuracy by individual audit parameters. Identifies which specific checkpoints are causing failures and need training or process improvement.')
    
    add_heading(doc, '6.1 Parameter Table Structure', level=2)
    
    add_heading(doc, 'Column: Parameter Name', level=3)
    doc.add_paragraph('• Source: Parameter Name field from Sheet1', style='List Bullet')
    doc.add_paragraph('• Examples: "Resume Format", "Experience Verification", "Background Check"', style='List Bullet')
    
    add_heading(doc, 'Column: Type (Critical/Non-Critical)', level=3)
    doc.add_paragraph('• Source: Parameter Type field', style='List Bullet')
    doc.add_paragraph('• Display: Badge showing "Critical" (red) or "Non-Critical" (blue)', style='List Bullet')
    doc.add_paragraph('• Purpose: Indicates parameter importance level', style='List Bullet')
    
    add_heading(doc, 'Column: Accuracy %', level=3)
    doc.add_paragraph('• Calculation: (Parameter Pass / (Parameter Total - Parameter NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Success rate for this specific parameter', style='List Bullet')
    doc.add_paragraph('• Color Coding: Applied based on critical vs non-critical thresholds', style='List Bullet')
    
    add_heading(doc, 'Columns: Pass, Fail, NA Counts', level=3)
    doc.add_paragraph('• Calculation: COUNT of each result type for the parameter', style='List Bullet')
    doc.add_paragraph('• Purpose: Absolute numbers for detailed analysis', style='List Bullet')
    
    add_heading(doc, '6.2 Usage Guidelines', level=2)
    doc.add_paragraph('• Sort by Accuracy: Click column to find lowest-performing parameters', style='List Bullet')
    doc.add_paragraph('• Filter by Type: Focus on critical or non-critical parameters', style='List Bullet')
    doc.add_paragraph('• Identify Training Needs: Parameters with <90% accuracy need attention', style='List Bullet')
    doc.add_paragraph('• Monthly Review: Compare parameter performance month-over-month', style='List Bullet')
    
    # ============================================================
    # RECRUITER INSIGHTS
    # ============================================================
    add_heading(doc, '7. Recruiter Insights Tab', level=1)
    
    doc.add_paragraph('Individual recruiter performance metrics. Essential for performance reviews, coaching, and identifying training needs at the individual level.')
    
    add_heading(doc, '7.1 Recruiter Performance Table', level=2)
    
    add_heading(doc, 'Column: Recruiter Name', level=3)
    doc.add_paragraph('• Source: Recruiter Name from Sheet1', style='List Bullet')
    doc.add_paragraph('• Display: Full name of individual recruiter', style='List Bullet')
    doc.add_paragraph('• Sortable: Alphabetically by name', style='List Bullet')
    
    add_heading(doc, 'Column: Accuracy %', level=3)
    doc.add_paragraph('• Calculation: (Recruiter Pass / (Recruiter Total - Recruiter NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Individual accuracy score', style='List Bullet')
    doc.add_paragraph('• Color Coding: Green ≥95%, Amber 90-94%, Red <90%', style='List Bullet')
    doc.add_paragraph('• Use: Performance evaluation and comparison', style='List Bullet')
    
    add_heading(doc, 'Column: Audit Count', level=3)
    doc.add_paragraph('• Calculation: COUNT(Audits) for recruiter', style='List Bullet')
    doc.add_paragraph('• Purpose: Volume of work audited', style='List Bullet')
    doc.add_paragraph('• Importance: Higher count = more reliable accuracy metric', style='List Bullet')
    
    add_heading(doc, 'Column: Error Count', level=3)
    doc.add_paragraph('• Calculation: COUNT(Fail results) for recruiter', style='List Bullet')
    doc.add_paragraph('• Purpose: Absolute number of errors made', style='List Bullet')
    doc.add_paragraph('• Use: Identify specific coaching opportunities', style='List Bullet')
    
    add_heading(doc, 'Column: Error %', level=3)
    doc.add_paragraph('• Calculation: (Error Count / (Total - NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Purpose: Error rate percentage', style='List Bullet')
    doc.add_paragraph('• Target: ≤5%', style='List Bullet')
    
    add_heading(doc, '7.2 Advanced Filters', level=2)
    doc.add_paragraph('• Region Filter: View recruiters by region', style='List Bullet')
    doc.add_paragraph('• Account Filter: Filter by client account', style='List Bullet')
    doc.add_paragraph('• Stage Filter: Focus on specific recruitment stages', style='List Bullet')
    doc.add_paragraph('• Time Period: Select month or date range', style='List Bullet')
    
    add_heading(doc, '7.3 Use Cases', level=2)
    doc.add_paragraph('• Performance Reviews: Export data for annual/quarterly reviews', style='List Bullet')
    doc.add_paragraph('• Coaching Plans: Identify low-performers needing support', style='List Bullet')
    doc.add_paragraph('• Best Practice Sharing: Identify top performers for knowledge sharing', style='List Bullet')
    doc.add_paragraph('• Training ROI: Track accuracy improvement after training', style='List Bullet')
    
    # ============================================================
    # STRATEGIC OVERVIEW
    # ============================================================
    add_heading(doc, '8. Strategic Overview Tab', level=1)
    
    doc.add_paragraph('Leadership-level view showing regional and head-wise performance with interactive India map. Used for strategic planning and resource allocation.')
    
    add_heading(doc, '8.1 Regional Head Performance Table', level=2)
    
    add_heading(doc, 'Table Structure', level=3)
    doc.add_paragraph('• Column 1: Regional Head Name (from Regional Head field)', style='List Bullet')
    doc.add_paragraph('• Column 2: Region (North/South/East/West)', style='List Bullet')
    doc.add_paragraph('• Column 3: Accuracy % - calculated per head', style='List Bullet')
    doc.add_paragraph('• Column 4: Audit Count - total audits under this head', style='List Bullet')
    doc.add_paragraph('• Column 5: Population - total population under this head', style='List Bullet')
    doc.add_paragraph('• Column 6: Pass/Fail/NA - breakdown of results', style='List Bullet')
    
    add_heading(doc, 'Calculations', level=3)
    doc.add_paragraph('• Accuracy = (Head Pass / (Head Total - Head NA)) × 100', style='List Bullet')
    doc.add_paragraph('• Sample % = (Head Audits / Head Population) × 100', style='List Bullet')
    doc.add_paragraph('• All metrics aggregated by Regional Head dimension', style='List Bullet')
    
    add_heading(doc, '8.2 Practice Head Performance Table', level=2)
    doc.add_paragraph('Same structure as Regional Head table but grouped by Practice Head', style='List Bullet')
    doc.add_paragraph('Practice areas: Technology, Finance, Healthcare, Manufacturing, etc.', style='List Bullet')
    
    add_heading(doc, '8.3 Interactive India Map', level=2)
    
    add_heading(doc, 'Map Regions', level=3)
    doc.add_paragraph('• North Region: Delhi, Chandigarh, Jaipur, etc.', style='List Bullet')
    doc.add_paragraph('• South Region: Bangalore, Hyderabad, Chennai, etc.', style='List Bullet')
    doc.add_paragraph('• East Region: Kolkata, Bhubaneswar, Patna, etc.', style='List Bullet')
    doc.add_paragraph('• West Region: Mumbai, Pune, Ahmedabad, etc.', style='List Bullet')
    
    add_heading(doc, 'Color Coding on Map', level=3)
    doc.add_paragraph('• Green regions: Accuracy ≥95%', style='List Bullet')
    doc.add_paragraph('• Amber regions: Accuracy 90-94%', style='List Bullet')
    doc.add_paragraph('• Red regions: Accuracy <90%', style='List Bullet')
    
    add_heading(doc, 'Interactive Features', level=3)
    doc.add_paragraph('• Hover: Shows tooltip with region name and accuracy %', style='List Bullet')
    doc.add_paragraph('• Click: Opens detailed drill-down with city-wise breakdown', style='List Bullet')
    doc.add_paragraph('• Zoom: Click zoom buttons to see specific areas', style='List Bullet')
    
    # ============================================================
    # PAGE 8: PROJECT, RCA & CAPA
    # ============================================================
    doc.add_page_break()
    add_heading(doc, '9. Project Summary Tab', level=1)
    
    doc.add_paragraph('Project-wise performance metrics for tracking individual projects/engagements.')
    
    add_heading(doc, '9.1 Project Table Columns', level=2)
    doc.add_paragraph('• Project Name: Name of project/engagement', style='List Bullet')
    doc.add_paragraph('• Account: Client account for the project', style='List Bullet')
    doc.add_paragraph('• Region: Geographic region', style='List Bullet')
    doc.add_paragraph('• Accuracy %: Project-specific accuracy', style='List Bullet')
    doc.add_paragraph('• Audit Count: Number of audits for project', style='List Bullet')
    doc.add_paragraph('• Population: Total population for project', style='List Bullet')
    doc.add_paragraph('• Pass/Fail/NA: Result breakdown', style='List Bullet')
    
    add_heading(doc, '9.2 Filters', level=2)
    doc.add_paragraph('• Region: Filter by geographic region', style='List Bullet')
    doc.add_paragraph('• Account: Filter by client account', style='List Bullet')
    doc.add_paragraph('• Month: Filter by time period', style='List Bullet')
    doc.add_paragraph('• Status: Filter by project status (Active/Completed)', style='List Bullet')
    
    # ============================================================
    # RCA & CAPA
    # ============================================================
    add_heading(doc, '10. RCA & CAPA Tab', level=1)
    
    doc.add_paragraph('Root Cause Analysis and Corrective & Preventive Actions. Critical for quality improvement initiatives.')
    
    add_heading(doc, '10.1 Top 5 Error Categories', level=2)
    
    add_heading(doc, 'Display', level=3)
    doc.add_paragraph('• Bar chart showing 5 most common error types', style='List Bullet')
    doc.add_paragraph('• X-axis: Error category name', style='List Bullet')
    doc.add_paragraph('• Y-axis: Count of occurrences', style='List Bullet')
    doc.add_paragraph('• Sorted: Highest to lowest by count', style='List Bullet')
    
    add_heading(doc, 'Calculation', level=3)
    doc.add_paragraph('• Source: "Reason for Failure" field from Sheet1', style='List Bullet')
    doc.add_paragraph('• Grouping: COUNT failures grouped by reason', style='List Bullet')
    doc.add_paragraph('• Top 5: SELECT TOP 5 ORDER BY Count DESC', style='List Bullet')
    
    add_heading(doc, '10.2 Reason-wise Breakdown Table', level=2)
    doc.add_paragraph('• Column 1: Failure Reason (e.g., "Incomplete Documentation")', style='List Bullet')
    doc.add_paragraph('• Column 2: Count (number of occurrences)', style='List Bullet')
    doc.add_paragraph('• Column 3: Percentage (% of total failures)', style='List Bullet')
    doc.add_paragraph('• Column 4: Trend (↑ increasing, ↓ decreasing, → stable)', style='List Bullet')
    
    add_heading(doc, '10.3 Action Items Section', level=2)
    doc.add_paragraph('• CAPA ID: Unique identifier for action item', style='List Bullet')
    doc.add_paragraph('• Issue Description: Detailed problem statement', style='List Bullet')
    doc.add_paragraph('• Root Cause: Identified root cause using 5-Why analysis', style='List Bullet')
    doc.add_paragraph('• Corrective Action: Immediate fix implemented', style='List Bullet')
    doc.add_paragraph('• Preventive Action: Long-term solution to prevent recurrence', style='List Bullet')
    doc.add_paragraph('• Owner: Person responsible for action', style='List Bullet')
    doc.add_paragraph('• Due Date: Target completion date', style='List Bullet')
    doc.add_paragraph('• Status: Open/In Progress/Completed/Closed', style='List Bullet')
    
    add_heading(doc, '10.4 How to Use RCA & CAPA', level=2)
    doc.add_paragraph('1. Identify top error categories from chart', style='List Number')
    doc.add_paragraph('2. Review detailed breakdown table', style='List Number')
    doc.add_paragraph('3. Create action items for each major issue', style='List Number')
    doc.add_paragraph('4. Assign owners and set due dates', style='List Number')
    doc.add_paragraph('5. Track status weekly until completion', style='List Number')
    doc.add_paragraph('6. Verify effectiveness after implementation', style='List Number')
    
    # ============================================================
    # PAGE 9-10: CSAT, COMPETITIVE INTELLIGENCE, SLA
    # ============================================================
    doc.add_page_break()
    add_heading(doc, '11. CSAT Tab - Customer Satisfaction', level=1)
    
    doc.add_paragraph('Measures client satisfaction through survey responses. Scale: 1-7 (1=Very Dissatisfied, 7=Very Satisfied).')
    
    add_heading(doc, '11.1 Top Box Cards', level=2)
    
    add_heading(doc, 'Card 1: Overall CSAT Score', level=3)
    doc.add_paragraph('• Calculation: Average of all ratings (Sum / Count)', style='List Bullet')
    doc.add_paragraph('• Current: 5.13 (down from 5.95 previous cycle)', style='List Bullet')
    doc.add_paragraph('• Target: ≥6.0', style='List Bullet')
    doc.add_paragraph('• Color: Green ≥6.0, Amber 5.0-5.9, Red <5.0', style='List Bullet')
    
    add_heading(doc, 'Card 2: Top Box % (Promoters)', level=3)
    doc.add_paragraph('• Calculation: (Ratings 6-7 / Total) × 100', style='List Bullet')
    doc.add_paragraph('• Meaning: Percentage of satisfied customers', style='List Bullet')
    doc.add_paragraph('• Target: ≥70%', style='List Bullet')
    
    add_heading(doc, 'Card 3: Response Rate', level=3)
    doc.add_paragraph('• Calculation: (Responses / Surveys Sent) × 100', style='List Bullet')
    doc.add_paragraph('• Current: 61% (absolute numbers increasing)', style='List Bullet')
    doc.add_paragraph('• Target: ≥65%', style='List Bullet')
    
    add_heading(doc, '11.2 Rating Distribution Chart', level=2)
    doc.add_paragraph('• Bar chart: X-axis = Ratings 1-7, Y-axis = Count', style='List Bullet')
    doc.add_paragraph('• Color coded: Ratings 1-3 (Red), 4-5 (Amber), 6-7 (Green)', style='List Bullet')
    doc.add_paragraph('• Shows: Distribution pattern of customer satisfaction', style='List Bullet')
    
    add_heading(doc, '11.3 Account-wise CSAT Table', level=2)
    doc.add_paragraph('• Columns: Account Name, Avg Score, Promoters %, Response Count', style='List Bullet')
    doc.add_paragraph('• Sorted: By average score (lowest to highest) to identify problem accounts', style='List Bullet')
    doc.add_paragraph('• Filterable: By account type, region, time period', style='List Bullet')
    
    add_heading(doc, '11.4 Key Insights & Analysis', level=2)
    doc.add_paragraph('Current Key Findings:', style='List Bullet')
    doc.add_paragraph('  - Response Rate: 61% (absolute numbers up, but percentage target not met)', style='List Bullet 2')
    doc.add_paragraph('  - Overall Satisfaction: 5.13 (dropped from 5.95)', style='List Bullet 2')
    doc.add_paragraph('  - Most Impact: OPS level and Non-Key accounts', style='List Bullet 2')
    doc.add_paragraph('Positive Feedback:', style='List Bullet')
    doc.add_paragraph('  - End-to-end hiring process appreciated', style='List Bullet 2')
    doc.add_paragraph('  - Strong engagement and flexibility', style='List Bullet 2')
    doc.add_paragraph('  - Team commitment highly rated', style='List Bullet 2')
    doc.add_paragraph('Focus Areas for Improvement:', style='List Bullet')
    doc.add_paragraph('  - Niche hiring quality needs enhancement', style='List Bullet 2')
    doc.add_paragraph('  - Overall team quality consistency', style='List Bullet 2')
    doc.add_paragraph('  - More actionable insights needed', style='List Bullet 2')
    doc.add_paragraph('  - Better candidate profiles required', style='List Bullet 2')
    doc.add_paragraph('  - Increased tech/IT focus needed', style='List Bullet 2')
    
    # ============================================================
    # COMPETITIVE INTELLIGENCE
    # ============================================================
    add_heading(doc, '12. Competitive Intelligence Tab', level=1)
    
    doc.add_paragraph('Industry benchmark comparison to understand Taggd performance relative to market standards.')
    
    add_heading(doc, '12.1 Benchmark Categories', level=2)
    doc.add_paragraph('• Global RPO: Global Recruitment Process Outsourcing standards', style='List Bullet')
    doc.add_paragraph('• Indian RPO: India-specific RPO benchmarks', style='List Bullet')
    doc.add_paragraph('• Global Service Provider: Worldwide service provider standards', style='List Bullet')
    doc.add_paragraph('• Indian Service Provider: India service provider benchmarks', style='List Bullet')
    
    add_heading(doc, '12.2 Comparison Metrics', level=2)
    
    add_heading(doc, 'Quality Score (%)', level=3)
    doc.add_paragraph('• Taggd Score vs Industry Average', style='List Bullet')
    doc.add_paragraph('• Calculation: Overall accuracy percentage', style='List Bullet')
    doc.add_paragraph('• Display: Side-by-side bar comparison', style='List Bullet')
    
    add_heading(doc, 'CSAT Score (1-7)', level=3)
    doc.add_paragraph('• Taggd CSAT vs Industry Benchmark', style='List Bullet')
    doc.add_paragraph('• Shows: Customer satisfaction comparison', style='List Bullet')
    
    add_heading(doc, 'TAT (Days)', level=3)
    doc.add_paragraph('• Turn Around Time comparison', style='List Bullet')
    doc.add_paragraph('• Taggd TAT vs Industry Standard', style='List Bullet')
    doc.add_paragraph('• Lower is better (faster delivery)', style='List Bullet')
    
    add_heading(doc, 'Offer-to-Join %', level=3)
    doc.add_paragraph('• Acceptance rate comparison', style='List Bullet')
    doc.add_paragraph('• (Joined / Offers Made) × 100', style='List Bullet')
    doc.add_paragraph('• Higher is better (better candidate quality)', style='List Bullet')
    
    add_heading(doc, '12.3 Filtering Options', level=2)
    doc.add_paragraph('• Switch between benchmark categories using filter buttons', style='List Bullet')
    doc.add_paragraph('• View all metrics simultaneously or focus on specific metric', style='List Bullet')
    doc.add_paragraph('• Export comparison data for presentations', style='List Bullet')
    
    # ============================================================
    # SLA OVERVIEW
    # ============================================================
    add_heading(doc, '13. SLA Overview Tab', level=1)
    
    doc.add_paragraph('Service Level Agreement compliance tracking. Monitors timely delivery across all stages.')
    
    add_heading(doc, '13.1 SLA Summary Cards', level=2)
    
    add_heading(doc, 'Card 1: Total SLA Met', level=3)
    doc.add_paragraph('• Calculation: COUNT(SLA Status = "Met")', style='List Bullet')
    doc.add_paragraph('• Shows: Number of cases delivered within SLA', style='List Bullet')
    doc.add_paragraph('• Display: Green card with checkmark icon', style='List Bullet')
    
    add_heading(doc, 'Card 2: Total SLA Breached', level=3)
    doc.add_paragraph('• Calculation: COUNT(SLA Status = "Breached")', style='List Bullet')
    doc.add_paragraph('• Shows: Number of cases delivered after SLA deadline', style='List Bullet')
    doc.add_paragraph('• Display: Red card with alert icon', style='List Bullet')
    
    add_heading(doc, 'Card 3: SLA Compliance %', level=3)
    doc.add_paragraph('• Calculation: (SLA Met / (SLA Met + SLA Breached)) × 100', style='List Bullet')
    doc.add_paragraph('• Target: ≥80% compliance', style='List Bullet')
    doc.add_paragraph('• Color: Green ≥80%, Amber 70-79%, Red <70%', style='List Bullet')
    
    add_heading(doc, 'Card 4: Average TAT (Days)', level=3)
    doc.add_paragraph('• Calculation: AVERAGE(Actual TAT Days)', style='List Bullet')
    doc.add_paragraph('• Shows: Average time taken across all stages', style='List Bullet')
    doc.add_paragraph('• Compared to: Target TAT for each stage', style='List Bullet')
    
    add_heading(doc, '13.2 Stage-wise SLA Table', level=2)
    doc.add_paragraph('• Column 1: Stage Name (Sourcing, Screening, Interview, Onboarding)', style='List Bullet')
    doc.add_paragraph('• Column 2: Target TAT (SLA deadline in days)', style='List Bullet')
    doc.add_paragraph('• Column 3: Actual Avg TAT (average days taken)', style='List Bullet')
    doc.add_paragraph('• Column 4: SLA Met Count', style='List Bullet')
    doc.add_paragraph('• Column 5: SLA Breached Count', style='List Bullet')
    doc.add_paragraph('• Column 6: Compliance % (color-coded)', style='List Bullet')
    
    add_heading(doc, '13.3 SLA Trend Chart', level=2)
    doc.add_paragraph('• Line chart: Shows monthly SLA compliance % trend', style='List Bullet')
    doc.add_paragraph('• Reference line: 80% target line', style='List Bullet')
    doc.add_paragraph('• Purpose: Track improvement or deterioration over time', style='List Bullet')
    
    add_heading(doc, '13.4 Filters', level=2)
    doc.add_paragraph('• Account: Filter by client account', style='List Bullet')
    doc.add_paragraph('• Region: Filter by geographic region', style='List Bullet')
    doc.add_paragraph('• Stage: Focus on specific stage', style='List Bullet')
    doc.add_paragraph('• Month: Select time period', style='List Bullet')
    
    # ============================================================
    # FINAL SECTION
    # ============================================================
    doc.add_page_break()
    add_heading(doc, '14. Data Upload Process', level=1)
    
    add_heading(doc, 'Step-by-Step Upload', level=2)
    doc.add_paragraph('1. Click "Upload Excel Data" button (top right of any tab)', style='List Number')
    doc.add_paragraph('2. Select Base File.xlsx from your computer', style='List Number')
    doc.add_paragraph('3. Wait for upload progress bar (5-10 seconds)', style='List Number')
    doc.add_paragraph('4. System validates file structure and required sheets', style='List Number')
    doc.add_paragraph('5. Success message appears when upload complete', style='List Number')
    doc.add_paragraph('6. Dashboard automatically refreshes with new data', style='List Number')
    doc.add_paragraph('7. All tabs updated simultaneously', style='List Number')
    
    add_heading(doc, 'Excel File Requirements', level=2)
    doc.add_paragraph('Required Sheets:', style='List Bullet')
    doc.add_paragraph('  1. Sheet1: Main audit data with all parameters', style='List Bullet 2')
    doc.add_paragraph('  2. CSAT: Customer satisfaction survey responses', style='List Bullet 2')
    doc.add_paragraph('  3. SLA Data: SLA tracking with TAT information', style='List Bullet 2')
    
    doc.add_paragraph('Required Columns in Sheet1:', style='List Bullet')
    doc.add_paragraph('  • Audit Date, Account Name, Region, Stage', style='List Bullet 2')
    doc.add_paragraph('  • Recruiter Name, Regional Head, Practice Head', style='List Bullet 2')
    doc.add_paragraph('  • Parameter Name, Critical/Non-Critical', style='List Bullet 2')
    doc.add_paragraph('  • Result (Pass/Fail/NA)', style='List Bullet 2')
    doc.add_paragraph('  • Total Population, Opportunity Count', style='List Bullet 2')
    
    # ============================================================
    # TROUBLESHOOTING & CONTACT
    # ============================================================
    add_heading(doc, '15. Troubleshooting Guide', level=1)
    
    issues_solutions = [
        ('Dashboard shows "No data available"', 'Upload Base File.xlsx using Upload button, ensure all required sheets present'),
        ('Filters not responding', 'Refresh browser page (F5), clear cache (Ctrl+Shift+Delete), try again'),
        ('Charts displaying incorrectly', 'Clear browser cache, try different browser (Chrome recommended), check zoom level'),
        ('Upload failing', 'Verify Excel format (.xlsx), check all required sheets exist, ensure no special characters in column names'),
        ('Wrong calculations displayed', 'Verify Excel column data types (numbers as numbers, not text), check for blank rows, ensure date formats correct'),
        ('Colors not showing', 'Check browser supports CSS3, disable browser extensions, try incognito mode'),
        ('Month-wise chart legend overlapping', 'This is fixed - legend now positioned above chart area'),
        ('Slow performance', 'Close other browser tabs, clear browser cache, check internet connection speed')
    ]
    
    for issue, solution in issues_solutions:
        p = doc.add_paragraph()
        run = p.add_run(f'Issue: {issue}')
        run.bold = True
        p = doc.add_paragraph(f'Solution: {solution}', style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
    
    # ============================================================
    # FOOTER
    # ============================================================
    doc.add_paragraph('_' * 100)
    doc.add_paragraph()
    
    # Contact table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light List Accent 1'
    
    contact_data = [
        ('Dashboard URL', 'https://3000-iiqzlm50w1twpgggr1553-82b888ba.sandbox.novita.ai'),
        ('Password', 'Excellence@2026'),
        ('Support Team', 'Business Excellence Team'),
        ('Document Version', '1.0'),
        ('Last Updated', 'March 2026')
    ]
    
    for i, (key, value) in enumerate(contact_data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        # Bold the key cell
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('End of Document').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    output_path = '/home/user/webapp/public/downloads/Quality_Dashboard_Complete_10Page_SOP.docx'
    doc.save(output_path)
    print(f"✅ Detailed 10-page SOP created: {output_path}")
    print(f"📄 Covers: All 12 tabs, every card, every view, all calculations")
    print(f"📊 Includes: Formulas, examples, color coding, filters, troubleshooting")
    
    return output_path

if __name__ == '__main__':
    create_detailed_10page_sop()
